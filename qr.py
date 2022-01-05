import shutil
from pathlib import Path
from docx import Document   # pip install python-docx
from docx.shared import Cm
import cv2 as cv            # pip install opencv-python
import docx2pdf

import qrcode


TEMP = './temp'
TEMP_WORD = './temp_word'


def _create_directory(directory):
    Path(directory).mkdir(parents=True, exist_ok=True)


def get_data(filename) -> str:
    image = cv.imread(filename)
    detector = cv.QRCodeDetector()

    data, _, _ = detector.detectAndDecode(image)
    return data


class Code:
    def __init__(self, data):
        self._data = data

    def _save_image(self, filename, ending):
        _qr = qrcode.QRCode()
        _qr.add_data(self._data)
        _qr.make(fit=True)
        _img = _qr.make_image(fill_color="black", back_color="white")
        _img.save(f'{filename}.{ending}')

    def save_jpg(self, filename: str):
        self._save_image(filename, 'jpg')

    def save_png(self, filename: str):
        self._save_image(filename, 'png')

    def save_docx(self, filename: str):
        document = Document()
        paragraph = document.add_paragraph()
        runner = paragraph.add_run()
        _create_directory(TEMP)
        for index, date in enumerate(self._data):
            single_qr = Code(date)
            single_qr.save_jpg(f'{TEMP}/{index}')
            runner.add_picture(f'{TEMP}/{index}.jpg', width=Cm(5))
        document.save(f'./{filename}.docx')
        shutil.rmtree(TEMP)

    def save_pdf(self, filename: str):
        _create_directory(TEMP_WORD)
        self.save_docx(f'{TEMP_WORD}/{filename}')
        docx2pdf.convert(f'{TEMP_WORD}/{filename}.docx', './')
        shutil.rmtree(TEMP_WORD)
